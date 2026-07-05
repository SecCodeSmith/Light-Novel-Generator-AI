"""Novel exports: chapters as a ZIP archive, full story as PDF or Word (.docx).

All builders return raw bytes ready to stream as a download. Chapters that are
not written yet are exported with their planned summary as a placeholder so a
"full story" document always covers the whole outline.
"""
import io
import re
import zipfile
from pathlib import Path

from . import graph

# Unicode TTF candidates: Debian container first (fonts-dejavu-core, installed in
# the Dockerfile), Windows fallbacks for running the backend outside Docker.
_FONT_CANDIDATES = [
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
]


def safe_filename(title: str) -> str:
    name = re.sub(r"[^\w \-]+", "", title or "", flags=re.UNICODE).strip()
    name = re.sub(r"\s+", " ", name)
    return name or "story"


def _load_story(story_id: str) -> tuple[dict, list[dict]]:
    story = graph.get_story(story_id) or {}
    chapters = []
    for ch in graph.get_chapters(story_id):
        full = graph.get_chapter(story_id, ch["number"]) or ch
        chapters.append(full)
    return story, chapters


def _chapter_heading(ch: dict) -> str:
    title = ch.get("title") or f"Chapter {ch['number']}"
    return f"Chapter {ch['number']}: {title}" if title != f"Chapter {ch['number']}" else title


def _chapter_body(ch: dict) -> tuple[str, bool]:
    """Return (text, written). Placeholder for chapters that have no prose yet."""
    text = ch.get("text")
    if text:
        return text, True
    return "(This chapter is not written yet.)\n\nPlanned summary:\n" + (ch.get("summary") or "—"), False


def _outline_markdown(story: dict, chapters: list[dict]) -> str:
    lines = [f"# {story.get('title', 'Untitled')} — outline\n"]
    if story.get("premise"):
        lines.append(f"**Premise:** {story['premise']}\n")
    if story.get("style"):
        lines.append(f"**Style:** {story['style']}\n")
    for ch in chapters:
        lines.append(f"## {_chapter_heading(ch)} [{ch.get('status', 'planned')}]\n")
        lines.append((ch.get("summary") or "") + "\n")
    return "\n".join(lines)


# --------------------------------------------------------------------- ZIP

def chapters_zip(story_id: str) -> bytes:
    story, chapters = _load_story(story_id)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("00 - Outline.md", _outline_markdown(story, chapters))
        for ch in chapters:
            if not ch.get("text"):
                continue  # the zip carries finished prose; the outline covers the rest
            name = f"{ch['number']:02d} - {safe_filename(ch.get('title') or 'chapter')}.md"
            zf.writestr(name, f"# {_chapter_heading(ch)}\n\n{ch['text']}\n")
    return buf.getvalue()


# --------------------------------------------------------------------- PDF

def _find_fonts() -> tuple[str, str] | None:
    for regular, bold in _FONT_CANDIDATES:
        if Path(regular).is_file() and Path(bold).is_file():
            return regular, bold
    return None


def story_pdf(story_id: str) -> bytes:
    from fpdf import FPDF

    story, chapters = _load_story(story_id)

    class _PDF(FPDF):
        family = "helvetica"

        def footer(self):
            self.set_y(-15)
            self.set_font(self.family, size=9)
            self.set_text_color(130, 130, 130)
            self.cell(0, 10, str(self.page_no()), align="C")

    pdf = _PDF(format="A4")
    fonts = _find_fonts()
    if fonts:
        pdf.add_font("Body", "", fonts[0])
        pdf.add_font("Body", "B", fonts[1])
        pdf.family = "Body"
        clean = str
    else:
        # Core fonts are latin-1 only; degrade gracefully instead of crashing.
        def clean(s: str) -> str:
            return s.encode("latin-1", "replace").decode("latin-1")

    pdf.set_auto_page_break(True, margin=18)

    # Title page
    pdf.add_page()
    pdf.set_font(pdf.family, "B", 28)
    pdf.set_text_color(20, 20, 20)
    pdf.ln(60)
    pdf.multi_cell(0, 14, clean(story.get("title", "Untitled")), align="C")
    if story.get("premise"):
        pdf.ln(10)
        pdf.set_font(pdf.family, "", 12)
        pdf.set_text_color(90, 90, 90)
        pdf.multi_cell(0, 7, clean(story["premise"]), align="C")

    for ch in chapters:
        body, _written = _chapter_body(ch)
        pdf.add_page()
        pdf.set_font(pdf.family, "B", 16)
        pdf.set_text_color(20, 20, 20)
        pdf.multi_cell(0, 9, clean(_chapter_heading(ch)))
        pdf.ln(4)
        pdf.set_font(pdf.family, "", 11)
        pdf.set_text_color(35, 35, 35)
        pdf.multi_cell(0, 6.5, clean(body))

    return bytes(pdf.output())


# -------------------------------------------------------------------- DOCX

def story_docx(story_id: str) -> bytes:
    import docx

    story, chapters = _load_story(story_id)
    doc = docx.Document()

    doc.add_heading(story.get("title", "Untitled"), 0)
    if story.get("premise"):
        p = doc.add_paragraph()
        run = p.add_run(story["premise"])
        run.italic = True

    for ch in chapters:
        doc.add_page_break()
        doc.add_heading(_chapter_heading(ch), level=1)
        body, _written = _chapter_body(ch)
        for para in body.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
